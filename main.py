import os
import subprocess
import tkinter as tk
from docx import Document
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.common.exceptions import NoSuchElementException, ElementNotInteractableException
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.action_chains import ActionChains
import pyautogui
import time
from docx import Document
import pyperclip
from tkinter import scrolledtext

def open_webpage(type):

    #檔案儲存位置
    folder_path = "C:/Documents/novel-translate"
    os.makedirs(folder_path, exist_ok=True)
    # 取得URL
    url = entry.get()
    
    # 創建瀏覽器驅動程式
    driver = webdriver.Chrome()
    
    # 打開網頁
    driver.get(url)
    
    # 使用Selenium尋找特定元素
    selected_text = ""
    
    # 迴圈遍歷每個元素
    for i in range(1, 10000):
        # 使用元素的 ID 值來定位元素
        element_id = "L" + str(i)
        try:
            element = driver.find_element(By.ID, element_id)
            selected_text += element.text + "\n"
        except NoSuchElementException:
            break
    
    # 創建 Word 文件
    doc = Document()
    doc.add_paragraph(selected_text)
    
    # 取得使用者輸入的保存名稱
    save_name = entry2.get()
    # 構建完整的文件路徑
    file_name = save_name + ".docx"
    file_path = os.path.join(folder_path, file_name)
    # 儲存 Word 文件
    doc.save(file_path)
    # 關閉瀏覽器
    driver.quit()
    # 關閉視窗
    window.destroy()
    #翻譯
    if type == "1":
        translate(file_path, folder_path)
    else:
        translateAllInOnce(file_path, folder_path)
    
def translate(file_path, folder_path):
    # 打開chrome
    bat_file_path = 'start_chrome_debug.bat'
    subprocess.call(bat_file_path, shell=True)
    # 創建 ChromeOptions
    options = webdriver.ChromeOptions()
    options.add_experimental_option("debuggerAddress", "localhost:9212")
    # 創建瀏覽器驅動程式
    print("創建瀏覽器驅動程式")
    driver = webdriver.Chrome(options=options)
    # 移動到網站
    print("移動到網站")
    driver.get("https://chat.forefront.ai")
    print(driver.title)
    #等待網站開好
    print("等")
    wait = WebDriverWait(driver, 10)
    wait.until(EC.presence_of_element_located((By.CLASS_NAME, "text-th-primary-dark")))
    time.sleep(3)
    # 讀取Word檔案
    print("讀取Word檔案")
    doc = Document(file_path)
    file_name = os.path.splitext(os.path.basename(file_path))[0]
    docA = Document()
    docB = Document()
    print("讀取Word檔案")
    # 儲存讀取到的文字
    print("儲存讀取到的文字")
    text = ""
    selected_text = ""
    # 讀取每個段落的內容
    print("讀取每個段落的內容")
    for paragraph in doc.paragraphs:
        lines = paragraph.text.split('\n')
        for line in lines:
            if line:
                print("for")
                # 將段落文字添加到text中
                text += line
                # 將文字添加到文檔中
                docA.add_paragraph(text + "\n")
                #添加翻譯要求
                print("翻譯成中文")
                text="翻譯成繁體中文\n" + text
                #將其複製到剪貼簿中
                pyperclip.copy(text)
                print(text)
                #再輸入框位置點左鍵
                pyautogui.click(x=2128, y=1972)
                pyautogui.hotkey('ctrl', 'v')
                time.sleep(1)
                pyautogui.press('enter')
                #等待答案
                time.sleep (5)
                #複製答案
                print("複製")
                # 移動到起始位置
                pyautogui.moveTo(x=1658, y=726)
                # 按住滑鼠左鍵
                pyautogui.mouseDown(button='left')
                # 移動到結束位置
                pyautogui.moveTo(x=3326, y=806, duration=0.5)
                # 釋放滑鼠左鍵
                pyautogui.mouseUp(button='left')
                # 複製
                pyautogui.hotkey('ctrl', 'c')
                # 貼上
                # 從剪貼簿中獲取文字
                selected_text = pyperclip.paste()
                print(selected_text)
                # 刪除特定文字
                selected_text = selected_text.replace("GPT-3.5", "")
                print(selected_text)
                # 將文字添加到文檔中
                docA.add_paragraph(selected_text + "\n")
                docB.add_paragraph(selected_text + "\n")
                # 清空text
                text = ""
                selected_text = ""
                time.sleep(1)
                # 移動到起始位置
                pyautogui.moveTo(x=470, y=774)
                pyautogui.click(button='left')
                time.sleep(0.5)
                pyautogui.moveTo(x=392, y=942)
                pyautogui.click(button='left')
                time.sleep(1)
                # 重新整理網頁
                pyautogui.press('f5')
                #等待網站開好
                print("等")
                wait = WebDriverWait(driver, 10)
                wait.until(EC.presence_of_element_located((By.CLASS_NAME, "text-th-primary-dark")))
                time.sleep(3)

    print("結束")
    # 關閉網站
    print("關閉")
    driver.quit()           
    #儲存翻譯
    print("存檔")
    file_nameB = file_name + "中文翻譯.docx"
    file_name = file_name + "中日對照.docx"
    file_path = os.path.join(folder_path, file_name)
    file_pathB = os.path.join(folder_path, file_nameB)
    print(file_path)
    print(file_pathB)
    # 儲存 Word 文件
    docA.save(file_path)
    docB.save(file_pathB)
    print("完成")
    # 關閉瀏覽器
    print("關閉")
    pyautogui.moveTo(x=3773, y=38)
    pyautogui.click(button='left')

def translateAllInOnce(file_path, folder_path):
    # 打開chrome
    bat_file_path = 'start_chrome_debug.bat'
    subprocess.call(bat_file_path, shell=True)
    # 創建 ChromeOptions
    options = webdriver.ChromeOptions()
    options.add_experimental_option("debuggerAddress", "localhost:9212")
    # 創建瀏覽器驅動程式
    print("創建瀏覽器驅動程式")
    driver = webdriver.Chrome(options=options)
    # 移動到網站
    print("移動到網站")
    driver.get("https://chat.forefront.ai")
    print(driver.title)
    #等待網站開好
    print("等")
    wait = WebDriverWait(driver, 10)
    wait.until(EC.presence_of_element_located((By.CLASS_NAME, "text-th-primary-dark")))
    time.sleep(3)
    # 讀取Word檔案
    print("讀取Word檔案")
    doc = Document(file_path)
    file_name = os.path.splitext(os.path.basename(file_path))[0]
    docA = Document()
    print("讀取Word檔案")
    # 儲存讀取到的文字
    print("儲存讀取到的文字")
    text = ""
    selected_text = ""
    # 將段落文字添加到text中
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    # 讀取每個段落的內容
    print("讀取每個段落的內容")
    #添加翻譯要求
    print("翻譯成中文")
    text="翻譯成繁體中文\n" + text
    #將其複製到剪貼簿中
    pyperclip.copy(text)
    print(text)
    #再輸入框位置點左鍵
    pyautogui.click(x=2128, y=1972)
    pyautogui.hotkey('ctrl', 'v')
    time.sleep(1)
    pyautogui.press('enter')
    #等待答案
    time.sleep(30)
    # 移動到起始位置
    pyautogui.moveTo(x = 2127, y = 968)
    # 按滑鼠左鍵
    pyautogui.click(button='left')
    # 移動到最下面
    pyautogui.press("end")
    time.sleep(1)
    # 移動到起始位置
    pyautogui.moveTo(x = 2717, y = 1701)
    # 按滑鼠左鍵
    pyautogui.click(button='left')
    # 從剪貼簿中獲取文字
    selected_text = pyperclip.paste()
    print(selected_text)
    ##### 刪除特定文字
    ####selected_text = selected_text.replace("GPT-3.5", "")
    ####print(selected_text)
    # 將文字添加到文檔中
    docA.add_paragraph(selected_text)
    # 清空text
    text = ""
    selected_text = ""
    time.sleep(1)           
    #儲存翻譯
    print("存檔")
    file_name = file_name + "中文翻譯.docx"
    file_pathA = os.path.join(folder_path, file_name)
    print(file_pathA)
    # 儲存 Word 文件
    docA.save(file_pathA)
    print("完成")
    checkanswer(file_path, folder_path, file_name)

def checkanswer(file_path, folder_path, file_name):
    # 创建新窗口
    window = tk.Tk()
    window.title("满意度调查")
    
    # 创建标签
    label = tk.Label(window, text="请问您对翻译是否满意？")
    label.pack()
    
    # 创建文本框
    text_box = scrolledtext.ScrolledText(window, height=10, width=50)
    text_box.pack()
    # 读取Word中的内容并显示在文本框中
    contentT = pyperclip.paste()
    text_box.insert(tk.END, contentT)
    
    # 创建按钮回调函数
    def yes_button_clicked():
        window.destroy()
        # 移動到起始位置
        pyautogui.moveTo(x=470, y=774)
        pyautogui.click(button='left')
        time.sleep(0.5)
        pyautogui.moveTo(x=392, y=942)
        pyautogui.click(button='left')
        time.sleep(1)
        # 關閉瀏覽器
        print("關閉")
        pyautogui.moveTo(x=3773, y=38)
        pyautogui.click(button='left')
        exit()
    
    def no_button_clicked():
        window.destroy()
        # 移動到起始位置
        pyautogui.moveTo(x=470, y=774)
        pyautogui.click(button='left')
        time.sleep(0.5)
        pyautogui.moveTo(x=392, y=942)
        pyautogui.click(button='left')
        time.sleep(1)
        # 關閉瀏覽器
        print("關閉")
        pyautogui.moveTo(x=3773, y=38)
        pyautogui.click(button='left')
        # 重新运行翻译
        translateAllInOnce(file_path, folder_path)
    
    def co_button_clicked(file_path, folder_path, file_name):
        window.destroy()
        continueTranslate(file_path, folder_path, file_name)

    # 创建按钮
    yes_button = tk.Button(window, text="满意", command=yes_button_clicked)
    yes_button.pack()
    
    no_button = tk.Button(window, text="不满意", command=no_button_clicked)
    no_button.pack()
    
    co_button = tk.Button(window, text="繼續翻譯", command=lambda: co_button_clicked (file_path, folder_path, file_name))
    co_button.pack()

    # 运行窗口的主循环
    window.mainloop()

def continueTranslate(file_path, folder_path, file_name):
    contentT = pyperclip.paste() + "\n"
    #再輸入框位置點左鍵
    pyautogui.click(x=2128, y=1972)
    Ctranslate = "請繼續翻譯"
    pyperclip.copy(Ctranslate)
    pyautogui.hotkey('ctrl', 'v')
    time.sleep(1)
    pyautogui.press('enter')
    #等待答案
    time.sleep(15)
    # 移動到起始位置
    pyautogui.moveTo(x = 2127, y = 968)
    # 按滑鼠左鍵
    pyautogui.click(button='left')
    # 移動到最下面
    pyautogui.press("end")
    time.sleep(1)
    # 移動到起始位置
    pyautogui.moveTo(x = 2717, y = 1701)
    # 按滑鼠左鍵
    pyautogui.click(button='left')
    # 從剪貼簿中獲取文字
    selected_text = pyperclip.paste()
    contentT += selected_text
    pyperclip.copy(contentT)
    print(contentT)
    ndoc = Document()
    #儲存翻譯
    ndoc.add_paragraph(contentT)
    file_path = os.path.join(folder_path, file_name)
    print(folder_path)
    print(file_path)
    # 儲存 Word 文件
    ndoc.save(file_path)
    checkanswer(file_path, folder_path, file_name)


# 創建視窗
window = tk.Tk()

# 建立標籤和輸入框
label = tk.Label(window, text="請輸入URL：")
label.pack()
entry = tk.Entry(window)
entry.pack()
label2 = tk.Label(window, text="請輸入保存名稱：")
label2.pack()
entry2 = tk.Entry(window)
entry2.pack()
# 创建按钮回调函数
def button1_clicked():
    open_webpage(1)

def button2_clicked():
    open_webpage(2)
# 建立按鈕
button1 = tk.Button(window, text="逐行翻譯(正確率較高)", command=button1_clicked)
button1.pack()
# 建立按鈕
button2 = tk.Button(window, text="整頁翻譯(效率較高)", command=button2_clicked)
button2.pack()

# 啟動主迴圈
window.mainloop()
